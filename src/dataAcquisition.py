import asyncio
import os
import io
from langchain_text_splitters import MarkdownHeaderTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers import EnsembleRetriever
import time, hashlib
from langchain_core.documents import Document
from src.table_to_json import TableToJSON
from src.univeral_table_parser import UniversalTableParser
from bs4 import BeautifulSoup
from src.llm_table_to_json import LLMTableToJson
from src.llm import llm
from src.image_analyzer import ImageAnalyzer
from langchain_google_genai import GoogleGenerativeAIEmbeddings

from src.rag_processor import RAGProcessor

import fitz
import pdfplumber
from docx import Document as DocxDocument


class DataAcquisition:

    def __init__(self):

        self.image_analyzer = ImageAnalyzer()
        self.embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-small-en",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

    def _get_session_db(self, session_id: str) -> Chroma:
            """In-memory Chroma collection per session — auto wiped when cleared."""
            return Chroma(
                collection_name=f"session_{session_id}",
                embedding_function=self.embeddings
            )

    def clear_session(self, session_id: str):
        """Wipes all chunks from session DB."""
        try:
            session_db = self._get_session_db(session_id)
            all_ids = session_db.get()["ids"]
            if all_ids:
                session_db.delete(ids=all_ids)
                print(f"[SESSION] Cleared {len(all_ids)} chunks for {session_id}")
            else:
                print(f"[SESSION] Nothing to clear for {session_id}")
        except Exception as e:
            print(f"[SESSION] Clear failed: {e}")

        try:
            crawl_db = self._get_session_db(f"crawl_{session_id}")
            ids = crawl_db.get()["ids"]
            if ids:
                crawl_db.delete(ids=ids)
                print(f"[CRAWL] Cleared {len(ids)} crawl chunks for {session_id}")
        except Exception as e:
            print(f"[CRAWL] Crawl clear failed: {e}")

    def extract_raw_tables(self,html):
        soup = BeautifulSoup(html, "html.parser")
        return [t.get_text(" ", strip=True) for t in soup.find_all("table")]

    def json_to_sentences(self,table_json, context=""):
        sentences = []

        if not isinstance(table_json, dict):
            return []

        for metric, values in table_json.items():
            if not isinstance(values, dict): continue
            for col, val in values.items():
                s = f"{metric} in {col} is {val}."
                if context:
                    s = f"{context}: {s}"
                sentences.append(s)
        return sentences

    def build_retriever(self, session_id: str = None):
        """Shared helper — builds EnsembleRetriever from current Chroma state."""

        all_docs = []

        session_docs = []
        if session_id:
            session_db = self._get_session_db(session_id)
            session_data = session_db.get(include=["documents", "metadatas"])
            if session_data["documents"]:
                session_docs = [
                    Document(page_content=d, metadata=m)
                    for d, m in zip(session_data["documents"], session_data["metadatas"])
                ]
                all_docs += session_docs

        print(f"[RETRIEVER]  Session: {len(session_docs)}")

        if not all_docs:
            empty_db = self._get_session_db("empty")
            return empty_db.as_retriever(search_type="mmr", search_kwargs={"k": 8})

        retrievers = []
        weights = []

        b_retriever = BM25Retriever.from_documents(all_docs)
        b_retriever.k = 10
        retrievers.append(b_retriever)
        weights.append(0.2)


        if session_docs:
            session_db = self._get_session_db(session_id)
            v_session = session_db.as_retriever(
                search_type="mmr",
                search_kwargs={"k": 15, "fetch_k": 40}
            )
            retrievers.append(v_session)
            weights.append(0.6)
        if len(retrievers) == 1:
            return retrievers[0]

        return EnsembleRetriever(retrievers=retrievers, weights=weights)

    def extract_text_from_pdf(self, doc_bytes: bytes) -> str:
        """Extract all text from PDF using pdfplumber."""
        text_parts = []
        with pdfplumber.open(io.BytesIO(doc_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()

                filtered_page = page
                if tables:
                    for table_obj in page.find_tables():
                        filtered_page = filtered_page.outside_bbox(table_obj.bbox)

                text = filtered_page.extract_text()
                if text:
                    text_parts.append(f"## Page {page_num}\n{text}")

                for table in tables:
                    if not table: continue
                    headers = table[0]
                    md_table = "| " + " | ".join(str(h or "") for h in headers) + " |\n"
                    md_table += "| " + " | ".join("---" for _ in headers) + " |\n"
                    for row in table[1:]:
                        md_table += "| " + " | ".join(str(c or "") for c in row) + " |\n"
                    text_parts.append(md_table)

        return "\n\n".join(text_parts)

    async def _extract_images_from_pdf(self, doc_bytes: bytes, filename: str) -> list[str]:
        """Extract embedded images from PDF and run VLM on each."""

        image_descriptions = []

        pdf = fitz.open(stream=doc_bytes, filetype="pdf")

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]

                    if len(image_bytes) < 5000:
                        continue

                    print(f"PDF Processing image {img_index + 1} on page {page_num + 1}...")
                    extracted_text = await self.image_analyzer.describe_image_with_vlm(image_bytes)

                    if extracted_text and "NO_TEXT_FOUND" not in extracted_text:
                        entry = (
                            f" Image from {filename} "
                            f"(Page {page_num + 1}, Image {img_index + 1}) ---\n"
                            f"{extracted_text}"
                        )
                        image_descriptions.append(entry)

                except Exception as e:
                    print(f"[PDF] Image extraction error: {e}")
                    continue

        pdf.close()
        return image_descriptions

    def extract_text_from_docx(self, doc_bytes: bytes) -> str:

        docx = DocxDocument(io.BytesIO(doc_bytes))
        parts = []

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text: continue
            style = para.style.name
            if "Heading 1" in style:
                parts.append(f"# {text}")
            elif "Heading 2" in style:
                parts.append(f"## {text}")
            elif "Heading 3" in style:
                parts.append(f"### {text}")
            else:
                parts.append(text)

        for table in docx.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows: continue
            md = "| " + " | ".join(rows[0]) + " |\n"
            md += "| " + " | ".join("---" for _ in rows[0]) + " |\n"
            for row in rows[1:]:
                md += "| " + " | ".join(row) + " |\n"
            parts.append(md)

        return "\n\n".join(parts)

    async def _extract_images_from_docx(self, doc_bytes: bytes, filename: str) -> list[str]:
        """Extract embedded images from DOCX and run VLM on each."""

        image_descriptions = []

        docx = DocxDocument(io.BytesIO(doc_bytes))

        for rel in docx.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                if len(image_bytes) < 5000:
                    continue

                print(f"[DOCX] Processing embedded image from {filename}...")
                extracted_text = await self.image_analyzer.describe_image_with_vlm(image_bytes)

                if extracted_text and "NO_TEXT_FOUND" not in extracted_text:
                    entry = (
                        f"--- Image from {filename} ---\n"
                        f"{extracted_text}"
                    )
                    image_descriptions.append(entry)

            except Exception as e:
                print(f"[DOCX] Image extraction error: {e}")
                continue

        return image_descriptions

    async def update_retriever_from_document_bytes(
            self, doc_bytes: bytes, mime_type: str, filename: str,session_id:str
    ):
        """
        PDF / DOCX / TXT → extract text + images → chunk → embed → retriever
        """
        rag_processor = RAGProcessor()

        text_content = ""
        image_descriptions = []

        if mime_type == "application/pdf":
            print(f"[DOC] Extracting text from PDF: {filename}")
            text_content = self.extract_text_from_pdf(doc_bytes)

            print(f"[DOC] Extracting images from PDF: {filename}")
            image_descriptions = await self._extract_images_from_pdf(doc_bytes, filename)

        elif "wordprocessingml" in mime_type:
            print(f"[DOC] Extracting text from DOCX: {filename}")
            text_content = self.extract_text_from_docx(doc_bytes)

            print(f"[DOC] Extracting images from DOCX: {filename}")
            image_descriptions = await self._extract_images_from_docx(doc_bytes, filename)

        elif mime_type == "text/plain":
            print(f"[DOC] Reading TXT: {filename}")
            text_content = doc_bytes.decode("utf-8", errors="ignore")

        else:
            print(f"[DOC] Unsupported mime type: {mime_type}")
            return self._get_session_db("empty").as_retriever()

        if not text_content and not image_descriptions:
            print("[DOC] Nothing extracted from document.")
            return self._get_session_db("empty").as_retriever()

        full_content = text_content
        if image_descriptions:
            full_content += "\n\n" + "\n\n".join(image_descriptions)

        print(f"[DOC] Total content length: {len(full_content)} chars")
        print(f"[DOC] Images extracted: {len(image_descriptions)}")

        source_id = f"{filename}_{int(time.time())}"
        chunks = rag_processor.chunk_content(
            content=full_content,
            source=source_id
        )

        session_db = self._get_session_db(session_id)
        session_db.add_documents(chunks)
        print(f"[DOC] Added {len(chunks)} chunks to session {session_id}")

        return self.build_retriever(session_id)

    async def process_webhook_data(self, markdown_content, html_content, url, image_scan=True, use_llm_tables=True):
        parser = UniversalTableParser()
        json_converter = TableToJSON()
        table_sentences = []
        image_descriptions = []

        if html_content:
            def parse_tables():
                tables = parser.parse_html(html_content)
                sentences = []
                for t in tables:
                    table_json = json_converter.convert(t)
                    sentences.extend(self.json_to_sentences(table_json, context=url))
                return sentences, tables

            print("[WEBHOOK] Starting table parse...")
            table_sentences, tables = await asyncio.to_thread(parse_tables)
            print(f"[WEBHOOK] Tables done: {len(table_sentences)} sentences, {len(tables)} tables")

            if len(table_sentences) < 3 and len(tables) > 0 and use_llm_tables:
                print("[WEBHOOK] Starting LLM table fallback...")
                raw_tables = self.extract_raw_tables(html_content)
                for raw in raw_tables:
                    try:
                        table_json = LLMTableToJson(llm, raw[:2000])
                        sentences = self.json_to_sentences(table_json, context=url)
                        table_sentences.extend(sentences)
                        print("[WEBHOOK] LLM fallback done")
                    except Exception as e:
                        print("LLM fallback failed:", e)

            if image_scan:
                image_descriptions = await self.image_analyzer.process_all_images(html_content, url)

        if not table_sentences and markdown_content:
            def parse_md_tables():
                tables = parser.parse_markdown(markdown_content)
                sentences = []
                for t in tables:
                    table_json = json_converter.convert(t)
                    sentences.extend(self.json_to_sentences(table_json, context=url))
                return sentences

            table_sentences = await asyncio.to_thread(parse_md_tables)

        table_sentences = list(set(table_sentences))
        image_descriptions = list(set(image_descriptions))

        enriched_content = (
                (markdown_content or "") + "\n" +
                "\n".join(table_sentences) + "\n" +
                "\n".join(image_descriptions)
        )

        def split_content():
            headers_to_split_on = [("#", "Header 1"), ("##", "Header 2")]
            header_splits = MarkdownHeaderTextSplitter(
                headers_to_split_on=headers_to_split_on
            ).split_text(enriched_content)

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200,
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            chunks = text_splitter.split_documents(header_splits)

            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "source": url,
                    "last_updated": int(time.time()),
                    "chunk_id": f"{hashlib.md5(url.encode()).hexdigest()}_{i}"
                })
            return chunks

        print("[WEBHOOK] Starting split...")
        result = await asyncio.to_thread(split_content)
        print(f"[WEBHOOK] Split done: {len(result)} chunks")
        return result

    async def update_and_get_retriever(self, chunks, url, session_id: str, progress_callback=None):
        session_db = self._get_session_db(session_id)

        # ✅ Run delete in thread
        existing = await asyncio.to_thread(session_db.get, where={"source": url})
        if existing["ids"]:
            await asyncio.to_thread(session_db.delete, ids=existing["ids"])

        if chunks:
            # ✅ Embed in batches
            batch_size = 100
            total = (len(chunks) + batch_size - 1) // batch_size
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                batch_num = i // batch_size + 1
                await asyncio.to_thread(session_db.add_documents, batch)
                print(f"[EMBED] Batch {batch_num}/{total}")
                if progress_callback:
                    await progress_callback(batch_num, total)  # ✅ UI heartbeat

        return await asyncio.to_thread(self.build_retriever, session_id)

    async def update_retriever_from_image_bytes(self, image_bytes,session_id:str, source="user_image"):

        rag_processor = RAGProcessor()

        extracted_text = await self.image_analyzer.describe_image_with_vlm(image_bytes)

        if not extracted_text or not extracted_text.strip():
            print("VLM processing failed or returned empty content")
            return "",self._get_session_db("empty").as_retriever()

        print("\n=== VLM OUTPUT ===\n", extracted_text)

        chunks = rag_processor.chunk_content(
            content=extracted_text,
            source=f"{source}_{time.time()}"
        )

        session_db = self._get_session_db(session_id)
        session_db.add_documents(chunks)
        print(f"[IMAGE] Added {len(chunks)} chunks to session {session_id}")
        return extracted_text,self.build_retriever(session_id)



if __name__=='__main__':
    da = DataAcquisition()
    test_chunk = Document(
        page_content="Virat Kohli is an Indian cricketer.",
        metadata={"source": "test_url"}
    )
    da.update_and_get_retriever([test_chunk], 'test_url')
